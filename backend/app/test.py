import os, re, json, requests, warnings
from dotenv import load_dotenv
import fitz
from docx import Document
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from typing import TypedDict, Optional
warnings.filterwarnings("ignore")

from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import PromptTemplate
from langgraph.graph import StateGraph, END

load_dotenv()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
ADZUNA_COUNTRY = os.getenv("ADZUNA_COUNTRY", "gb")

# ── NLTK ─────────────────────────────────────────────────────────────────────
for resource, pkg in [
    ("tokenizers/punkt", "punkt"),
    ("tokenizers/punkt_tab/english", "punkt_tab"),
    ("corpora/stopwords", "stopwords"),
]:
    try:
        nltk.data.find(resource)
    except LookupError:
        nltk.download(pkg, quiet=True)


# ════════════════════════════════════════════════════════════════════════════
# STATE — the object that flows through every node
# ════════════════════════════════════════════════════════════════════════════
class CareerState(TypedDict):
    # Inputs
    resume_path: str
    goal: str

    # Node outputs (populated as the graph runs)
    resume: Optional[dict]
    advice: Optional[dict]
    universities: Optional[list]
    jobs: Optional[list]
    errors: list[str]          # accumulates non-fatal errors


# ════════════════════════════════════════════════════════════════════════════
# DOMAIN DETECTION
# ════════════════════════════════════════════════════════════════════════════
DOMAIN_KEYWORDS = {
    "machine learning":     ["python","machine learning","data science","tensorflow","pytorch","hugging face","nlp","scikit"],
    "software engineering": ["c++","java","python","algorithms","api","fastapi","git","docker","kubernetes"],
    "mechanical":           ["solidworks","autocad","thermodynamics","cad","ansys","fluid mechanics","manufacturing"],
    "civil":                ["autocad","revit","estimation","surveying","structural","staad","construction"],
    "electrical":           ["circuit","arduino","embedded","matlab","power systems","pcb"],
    "electronics":          ["vlsi","verilog","fpga","microcontroller","signals","communication"],
    "management":           ["project management","excel","finance","operations","marketing","leadership","mba"],
    "design":               ["photoshop","illustrator","figma","ui/ux","animation"],
    "biotech":              ["molecular biology","genetics","biochemistry","cell culture","clinical"],
    "finance":              ["accounting","valuation","financial modeling","investment","statistics","equity"],
    "law":                  ["legal research","contracts","litigation","corporate law","moot court"],
    "environmental":        ["sustainability","climate","gis","remote sensing","carbon"],
}


# ════════════════════════════════════════════════════════════════════════════
# NODE 1 — parse_resume
# ════════════════════════════════════════════════════════════════════════════
def parse_resume_node(state: CareerState) -> CareerState:
    print("\n[Node 1/4] 📄 Parsing resume...")
    errors = state.get("errors", [])

    try:
        file_path = state["resume_path"]
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            doc = fitz.open(file_path)
            text = "".join(p.get_text() for p in doc)
            doc.close()
        elif ext in (".docx", ".doc"):
            d = Document(file_path)
            parts = [p.text for p in d.paragraphs]
            for t in d.tables:
                for row in t.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            text = "\n".join(parts)
        else:
            raise ValueError("Unsupported file type. Use PDF or DOCX.")

        # Domain detection
        tokens = word_tokenize(text.lower())
        stop = set(stopwords.words("english"))
        filtered = " ".join(w for w in tokens if w not in stop and len(w) > 2)
        scores = {
            domain: sum(1 for kw in kws if re.search(r"\b" + re.escape(kw) + r"\b", filtered, re.I))
            for domain, kws in DOMAIN_KEYWORDS.items()
        }
        best = max(scores, key=scores.get)
        domain = best if scores[best] >= 3 else "general"

        # Field extraction
        lines = text.strip().split("\n")
        name_match = re.search(r"^[A-Z][a-z]+(?: [A-Z][a-z]+){0,2}$", text, re.MULTILINE)
        name = name_match.group().strip() if name_match else (lines[0].strip() if lines else "Unknown")

        email_m = re.search(r"[\w.\-]+@[\w.\-]+\.\w+", text)
        phone_m = re.search(r"(\+?\d[\d\s\-]{8,}\d)", text)

        skills_sec = re.search(r"(?:skills?|technologies)[:\-]?\s*([\s\S]{30,400}?)(?:\n[A-Z]|\Z)", text, re.I)
        skills_text = skills_sec.group(1) if skills_sec else text[:500]
        skills = list({s.strip() for s in re.split(r"[,|\n•]", skills_text) if 2 < len(s.strip()) < 30})[:15]

        experience = re.findall(
            r"([a-zA-Z ]+?(?:engineer|developer|intern|analyst|researcher|student|manager)\s*(?:at [A-Za-z ]+)?)\s*(\d{4}[\s\-]*(?:\d{4}|present)?)",
            text, re.I
        )
        experience = [f"{r.strip()} ({p.strip()})" for r, p in experience[:6]]

        education = re.findall(
            r"(bachelor|master|btech|mtech|b\.?e|m\.?e|phd|bs|ms)\s+(?:of |in )?\s*([a-zA-Z ]{3,40})",
            text, re.I
        )
        education = [f"{d.title()} in {f.title()}" for d, f in education[:3]]

        resume = {
            "name": name,
            "email": email_m.group() if email_m else "Not found",
            "phone": phone_m.group().strip() if phone_m else "Not found",
            "domain": domain,
            "skills": skills,
            "experience": experience,
            "education": education,
            "raw_text": text[:3000],
        }
        print(f"    ✓ Name: {resume['name']}  |  Domain: {resume['domain'].title()}")
        return {**state, "resume": resume, "errors": errors}

    except Exception as e:
        errors.append(f"[parse_resume] {e}")
        print(f"    ✗ Error: {e}")
        return {**state, "resume": None, "errors": errors}


# ════════════════════════════════════════════════════════════════════════════
# NODE 2 — gemini_advice
# ════════════════════════════════════════════════════════════════════════════
ADVICE_TEMPLATE = """
You are an expert career counselor and university advisor.

Candidate profile:
- Name: {name}
- Detected Domain: {domain}
- Skills: {skills}
- Education: {education}
- Experience: {experience}
- Career Goal: {goal}

Respond ONLY in valid JSON (no markdown, no explanation, no trailing commas):
{{
  "career_paths": [
    {{"title":"...","why_suitable":"...","skills_to_build":["..."],"salary_range_inr":"..."}}
  ],
  "course_suggestions": [
    {{"course_name":"...","platform":"...","why_recommended":"...","estimated_duration":"...","free_or_paid":"..."}}
  ],
  "recommended_universities": [
    {{"name":"...","country":"...","why":"...","search_keyword":"..."}}
  ],
  "job_search_keywords": ["...", "...", "..."],
  "skill_gaps": ["...", "..."],
  "advice_summary": "..."
}}

Return exactly 3 career_paths, 3 course_suggestions, 3 recommended_universities, 3 job_search_keywords.
"""

def gemini_advice_node(state: CareerState) -> CareerState:
    print("\n[Node 2/4] 🤖 Generating AI advice via Gemini...")
    errors = state.get("errors", [])

    if not state.get("resume"):
        errors.append("[gemini_advice] Skipped — no resume data.")
        return {**state, "advice": None, "errors": errors}

    resume = state["resume"]
    try:
        llm = ChatGoogleGenerativeAI(
            model="gemini-2.5-flash",
            google_api_key=GEMINI_API_KEY,
            temperature=0.2,
        )
        prompt = PromptTemplate(
            input_variables=["name","domain","skills","education","experience","goal"],
            template=ADVICE_TEMPLATE,
        )
        chain = prompt | llm
        result = chain.invoke({
            "name":       resume["name"],
            "domain":     resume["domain"],
            "skills":     ", ".join(resume["skills"]),
            "education":  ", ".join(resume["education"]) or "Not specified",
            "experience": ", ".join(resume["experience"]) or "Fresher",
            "goal":       state["goal"],
        })
        raw = re.sub(r"```(?:json)?", "", result.content).strip().strip("`")
        advice = json.loads(raw)
        print("    ✓ Advice generated successfully.")
        return {**state, "advice": advice, "errors": errors}

    except json.JSONDecodeError:
        errors.append("[gemini_advice] JSON parse failed.")
        print("    ✗ Could not parse Gemini JSON.")
        return {**state, "advice": None, "errors": errors}
    except Exception as e:
        errors.append(f"[gemini_advice] {e}")
        print(f"    ✗ Gemini error: {e}")
        return {**state, "advice": None, "errors": errors}


# ════════════════════════════════════════════════════════════════════════════
# NODE 3 — fetch_universities  (Hipolabs, free, no key)
# ════════════════════════════════════════════════════════════════════════════
def fetch_universities_node(state: CareerState) -> CareerState:
    print("\n[Node 3/4] 🎓 Fetching universities (Hipolabs)...")
    errors = state.get("errors", [])
    advice = state.get("advice")

    if not advice:
        return {**state, "universities": [], "errors": errors}

    enriched = []
    for uni in advice.get("recommended_universities", []):
        # Use Gemini-provided search_keyword — much more reliable than parsing the full name
        keyword = uni.get("search_keyword") or uni.get("name", "")
        # Strip noise just in case
        for noise in ["university", "institute", "e.g.", "(iits)", "(iisc)", "(iiit)", "-"]:
            keyword = keyword.lower().replace(noise, "")
        keyword = " ".join(keyword.split()[:3]).strip()

        results = []
        try:
            res = requests.get(
                "http://universities.hipolabs.com/search",
                params={"name": keyword},
                timeout=8
            )
            if res.status_code == 200:
                results = res.json()[:3]
        except Exception as e:
            errors.append(f"[universities] {e}")

        # Pick best country match or first result
        country_hint = uni.get("country", "").lower()
        match = next(
            (u for u in results if country_hint in u.get("country", "").lower()),
            results[0] if results else None
        )

        enriched.append({
            **uni,
            "website": match["web_pages"][0] if match and match.get("web_pages") else "Search on QS Rankings"
        })
        status = f"→ {match['web_pages'][0]}" if match else "→ not found in Hipolabs"
        print(f"    {uni['name'][:50]}  {status}")

    return {**state, "universities": enriched, "errors": errors}


# ════════════════════════════════════════════════════════════════════════════
# NODE 4 — fetch_jobs  (Remotive + Arbeitnow, both free, no key)
# ════════════════════════════════════════════════════════════════════════════
def fetch_jobs_node(state: CareerState) -> CareerState:
    print("\n[Node 4/4] 💼 Fetching live jobs (Remotive + Arbeitnow)...")
    errors = state.get("errors", [])
    advice = state.get("advice")
    resume = state.get("resume", {})

    keywords = (advice or {}).get("job_search_keywords") or [resume.get("domain", ""), state["goal"]]
    query = " ".join(keywords[:2])
    jobs = []

    # Source A: Remotive
    try:
        res = requests.get(
            "https://remotive.com/api/remote-jobs",
            params={"search": query, "limit": 4},
            timeout=10
        )
        if res.status_code == 200:
            for j in res.json().get("jobs", []):
                jobs.append({
                    "title":    j.get("title"),
                    "company":  j.get("company_name", "N/A"),
                    "location": j.get("candidate_required_location", "Remote"),
                    "salary":   j.get("salary") or "Not disclosed",
                    "tags":     ", ".join(j.get("tags", [])[:4]),
                    "url":      j.get("url"),
                    "source":   "Remotive",
                })
            print(f"    ✓ Remotive: {len(jobs)} jobs found")
    except Exception as e:
        errors.append(f"[jobs/remotive] {e}")
        print(f"    ✗ Remotive: {e}")

    # Source B: Arbeitnow
    try:
        res = requests.get(
            "https://www.arbeitnow.com/api/job-board-api",
            params={"search": query},
            timeout=10
        )
        if res.status_code == 200:
            arb_jobs = res.json().get("data", [])[:4]
            for j in arb_jobs:
                jobs.append({
                    "title":    j.get("title"),
                    "company":  j.get("company_name", "N/A"),
                    "location": j.get("location", "N/A"),
                    "salary":   "Not disclosed",
                    "tags":     ", ".join(j.get("tags", [])[:4]),
                    "url":      j.get("url"),
                    "source":   "Arbeitnow",
                })
            print(f"    ✓ Arbeitnow: {len(arb_jobs)} jobs found")
    except Exception as e:
        errors.append(f"[jobs/arbeitnow] {e}")
        print(f"    ✗ Arbeitnow: {e}")

    # Deduplicate
    seen, unique = set(), []
    for j in jobs:
        key = (j["title"], j["company"])
        if key not in seen:
            seen.add(key)
            unique.append(j)

    return {**state, "jobs": unique[:6], "errors": errors}


# ════════════════════════════════════════════════════════════════════════════
# CONDITIONAL EDGE — abort if resume parsing failed
# ════════════════════════════════════════════════════════════════════════════
def should_continue(state: CareerState) -> str:
    if not state.get("resume"):
        print("\n❌ Aborting pipeline — resume could not be parsed.")
        return "abort"
    return "continue"


# ════════════════════════════════════════════════════════════════════════════
# BUILD THE GRAPH
# ════════════════════════════════════════════════════════════════════════════
def build_graph() -> any:
    graph = StateGraph(CareerState)

    graph.add_node("parse_resume",       parse_resume_node)
    graph.add_node("gemini_advice",      gemini_advice_node)
    graph.add_node("fetch_universities", fetch_universities_node)
    graph.add_node("fetch_jobs",         fetch_jobs_node)

    # Entry point
    graph.set_entry_point("parse_resume")

    # Conditional: only proceed if resume parsed OK
    graph.add_conditional_edges(
        "parse_resume",
        should_continue,
        {"continue": "gemini_advice", "abort": END}
    )

    # Gemini → then universities AND jobs run in parallel via two edges
    graph.add_edge("gemini_advice",      "fetch_universities")
    graph.add_edge("fetch_universities", "fetch_jobs")
    graph.add_edge("fetch_jobs",         END)

    return graph.compile()


# ════════════════════════════════════════════════════════════════════════════
# DISPLAY
# ════════════════════════════════════════════════════════════════════════════
def section(title):
    print(f"\n{'═'*60}\n  {title}\n{'═'*60}")

def display_results(state: CareerState):
    resume = state.get("resume", {})
    advice = state.get("advice", {})

    section("📄 RESUME SUMMARY")
    print(f"  Name     : {resume.get('name')}")
    print(f"  Email    : {resume.get('email')}")
    print(f"  Domain   : {resume.get('domain','').title()}")
    print(f"  Education: {', '.join(resume.get('education', [])) or 'Not detected'}")
    print(f"  Skills   : {', '.join(resume.get('skills', [])[:10]) or 'Not detected'}")

    if not advice:
        print("\n⚠️  No AI advice available.")
        return

    section("🚀 RECOMMENDED CAREER PATHS")
    for i, cp in enumerate(advice.get("career_paths", []), 1):
        print(f"\n  {i}. {cp['title']}")
        print(f"     Why       : {cp['why_suitable']}")
        print(f"     Build     : {', '.join(cp.get('skills_to_build', []))}")
        print(f"     Salary    : {cp.get('salary_range_inr')}")

    section("⚠️  SKILL GAPS")
    for gap in advice.get("skill_gaps", []):
        print(f"  • {gap}")

    section("📚 COURSE RECOMMENDATIONS")
    for i, c in enumerate(advice.get("course_suggestions", []), 1):
        print(f"\n  {i}. {c['course_name']}  [{c['platform']}]  |  {c.get('free_or_paid')}")
        print(f"     Why: {c['why_recommended']}")
        print(f"     Duration: {c['estimated_duration']}")

    section("🎓 UNIVERSITIES")
    for u in state.get("universities", []):
        print(f"\n  {u['name']} ({u['country']})")
        print(f"     Why: {u['why']}")
        print(f"     🌐 {u.get('website')}")

    section("💼 LIVE JOB LISTINGS")
    jobs = state.get("jobs", [])
    if jobs:
        for i, j in enumerate(jobs, 1):
            print(f"\n  {i}. {j['title']}  —  {j['company']}  [{j['source']}]")
            print(f"     Location : {j['location']}")
            print(f"     Salary   : {j['salary']}")
            if j.get("tags"):
                print(f"     Tags     : {j['tags']}")
            print(f"     Apply    : {j['url']}")
    else:
        keywords = advice.get("job_search_keywords", [])
        print(f"  No live listings found. Search: {' | '.join(keywords)}")
        print("  → LinkedIn / Naukri / Indeed / Wellfound")

    section("💡 ADVISOR SUMMARY")
    print(f"\n  {advice.get('advice_summary', '')}\n")

    if state.get("errors"):
        print("⚠️  Non-fatal errors during run:")
        for e in state["errors"]:
            print(f"   • {e}")


# ════════════════════════════════════════════════════════════════════════════
# MAIN
# ════════════════════════════════════════════════════════════════════════════
def main():
    print("\n╔══════════════════════════════════════════════════╗")
    print("║  AI Career Advisor  (LangGraph + Gemini)          ║")
    print("║  Nodes: Resume → Gemini → Universities → Jobs     ║")
    print("╚══════════════════════════════════════════════════╝")

    if not GEMINI_API_KEY:
        print("❌ GEMINI_API_KEY not set in .env. Exiting.")
        return

    resume_path = input("\nResume path (.pdf / .docx): ").strip()
    goal        = input("Career goal: ").strip()

    initial_state: CareerState = {
        "resume_path": resume_path,
        "goal":        goal,
        "resume":      None,
        "advice":      None,
        "universities": None,
        "jobs":        None,
        "errors":      [],
    }

    print("\n⏳ Running LangGraph pipeline...")
    pipeline = build_graph()
    final_state = pipeline.invoke(initial_state)

    display_results(final_state)


if __name__ == "__main__":
    main()