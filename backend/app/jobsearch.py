import os
import re
import json
import requests
import fitz  # PyMuPDF
from docx import Document
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
import nltk
import warnings
warnings.filterwarnings("ignore")
def _ensure_nltk_data():
    import nltk
    needed = [
        ("tokenizers/punkt", "punkt"),
        # some scripts or tokenizers might request 'punkt_tab' — try to ensure it too
        ("tokenizers/punkt_tab/english", "punkt_tab"),
        ("corpora/stopwords", "stopwords"),
        # add more if you use them: ("taggers/averaged_perceptron_tagger","averaged_perceptron_tagger")
    ]
    for resource_path, package_name in needed:
        try:
            nltk.data.find(resource_path)
        except LookupError:
            try:
                nltk.download(package_name)
            except Exception as e:
                # last-resort: print error but continue; caller can decide how to proceed
                print(f"[NLTK] download failed for {package_name}: {e}")

_ensure_nltk_data()
# Optional Hugging Face integration
try:
    from transformers import pipeline
    HF_AVAILABLE = True
except ImportError:
    print("Transformers not installed; fallback mode only.")
    HF_AVAILABLE = False

# Auto-download NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)
try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords', quiet=True)

# --- DOMAIN DEFINITIONS ---

DOMAIN_KEYWORDS = {
    "machine learning": ["python", "machine learning", "data science", "tensorflow", "pytorch", "hugging face", "flask", "nlp"],
    "software engineering": ["c", "c++", "java", "python", "algorithms", "data structures", "api", "fastapi", "git", "docker"],
    "mechanical": ["solidworks", "autocad", "thermodynamics", "cad", "cam", "ansys", "fluid mechanics", "manufacturing"],
    "civil": ["autocad", "revit", "estimation", "surveying", "structural analysis", "staad", "construction", "planning"],
    "electrical": ["circuit design", "arduino", "embedded systems", "matlab", "power systems", "pcb design"],
    "electronics": ["vlsi", "verilog", "embedded", "fpga", "arduino", "microcontroller", "signals", "communication"],
    "management": ["project management", "excel", "finance", "business analysis", "operations", "marketing", "leadership"],
    "design": ["photoshop", "illustrator", "figma", "ui/ux", "animation", "video editing"],
    "arts": ["creative writing", "photography", "music", "art history", "painting", "film production"],
    "biotech": ["molecular biology", "genetics", "biochemistry", "cell culture", "research", "clinical trials"],
    "finance": ["excel", "data analysis", "accounting", "valuation", "financial modeling", "investment", "statistics"],
    "law": ["legal research", "contracts", "litigation", "corporate law", "moot court"],
    "education": ["lesson planning", "pedagogy", "curriculum design", "training", "assessment"],
    "environmental": ["sustainability", "climate", "gis", "remote sensing", "environmental policy", "carbon accounting"]
}

# --- COURSE & JOB LIBRARIES ---

COURSE_LIBRARY = {
    "machine learning": [
        {"course_name": "Machine Learning Specialization", "platform": "Coursera", "why_recommended": "Strengthens fundamentals and practical ML skills.", "estimated_duration": "8 weeks"},
        {"course_name": "Deep Learning with PyTorch", "platform": "Udemy", "why_recommended": "Hands-on model implementation training.", "estimated_duration": "6 weeks"}
    ],
    "software engineering": [
        {"course_name": "System Design Fundamentals", "platform": "Udemy", "why_recommended": "Prepares for scalable backend design interviews.", "estimated_duration": "5 weeks"},
        {"course_name": "Advanced Python Programming", "platform": "Coursera", "why_recommended": "Covers modern frameworks and tools for production code.", "estimated_duration": "4 weeks"}
    ],
    "mechanical": [
        {"course_name": "CAD & SolidWorks Masterclass", "platform": "Udemy", "why_recommended": "Strengthens 3D modeling and mechanical drawing.", "estimated_duration": "4 weeks"},
        {"course_name": "Thermodynamics and Fluid Mechanics", "platform": "NPTEL", "why_recommended": "Refreshes key mechanical principles.", "estimated_duration": "6 weeks"}
    ],
    "civil": [
        {"course_name": "Revit Architecture and Structural Design", "platform": "Coursera", "why_recommended": "Builds design proficiency for modern civil projects.", "estimated_duration": "6 weeks"},
        {"course_name": "Construction Project Management", "platform": "edX", "why_recommended": "Learn planning, cost estimation, and risk control.", "estimated_duration": "5 weeks"}
    ],
    "management": [
        {"course_name": "Business Analytics", "platform": "Wharton via Coursera", "why_recommended": "Bridges business decisions and data insights.", "estimated_duration": "4 weeks"},
        {"course_name": "Project Management Professional (PMP) Prep", "platform": "Udemy", "why_recommended": "Equips you for managerial certification exams.", "estimated_duration": "6 weeks"}
    ],
    "environmental": [
        {"course_name": "Sustainability and Climate Policy", "platform": "edX", "why_recommended": "Develops understanding of global climate frameworks.", "estimated_duration": "5 weeks"},
        {"course_name": "GIS and Remote Sensing", "platform": "Coursera", "why_recommended": "Learn spatial data handling for environmental systems.", "estimated_duration": "6 weeks"}
    ],
    "general": [
        {"course_name": "Critical Thinking and Problem Solving", "platform": "Coursera", "why_recommended": "Enhances analytical and communication skills.", "estimated_duration": "3 weeks"}
    ]
}

JOB_LIBRARY = {
    "machine learning": [
        {"job_title": "ML Engineer", "company_type": "Tech/AI Startup", "required_skills": ["Python", "TensorFlow", "Pandas"], "why_suitable": "Applies your ML and data preprocessing experience."},
        {"job_title": "Data Analyst", "company_type": "Consulting Firm", "required_skills": ["Excel", "SQL", "Visualization"], "why_suitable": "Uses your analytical and problem-solving abilities."}
    ],
    "software engineering": [
        {"job_title": "Backend Developer", "company_type": "SaaS Company", "required_skills": ["API Development", "Databases", "Python/Java"], "why_suitable": "Builds on your coding and deployment experience."},
        {"job_title": "Full Stack Intern", "company_type": "Startup", "required_skills": ["React", "Node.js", "FastAPI"], "why_suitable": "Leverages your web development projects."}
    ],
    "mechanical": [
        {"job_title": "Design Engineer", "company_type": "Manufacturing Firm", "required_skills": ["AutoCAD", "SolidWorks"], "why_suitable": "Fits your mechanical design experience."},
        {"job_title": "Production Intern", "company_type": "Automotive Plant", "required_skills": ["Manufacturing", "Quality Control"], "why_suitable": "Relates to your workshop and CAD exposure."}
    ],
    "civil": [
        {"job_title": "Site Engineer", "company_type": "Construction Company", "required_skills": ["Surveying", "AutoCAD"], "why_suitable": "Matches your civil design background."},
        {"job_title": "Structural Designer", "company_type": "Infrastructure Firm", "required_skills": ["STAAD", "Revit"], "why_suitable": "Applies your structural modeling knowledge."}
    ],
    "management": [
        {"job_title": "Operations Analyst", "company_type": "Consulting", "required_skills": ["Excel", "Data Analysis"], "why_suitable": "Fits your strategic and analytical skills."},
        {"job_title": "Marketing Associate", "company_type": "Corporate", "required_skills": ["Communication", "Digital Tools"], "why_suitable": "Utilizes your managerial training."}
    ],
    "environmental": [
        {"job_title": "Sustainability Analyst", "company_type": "NGO/Think Tank", "required_skills": ["Carbon Accounting", "Policy Research"], "why_suitable": "Matches your sustainability expertise."},
        {"job_title": "GIS Technician", "company_type": "Research Institute", "required_skills": ["Remote Sensing", "Data Mapping"], "why_suitable": "Applies your GIS and environmental data skills."}
    ],
    "general": [
        {"job_title": "Research Assistant", "company_type": "Academic Institution", "required_skills": ["Writing", "Data Collection"], "why_suitable": "Good fit for multidisciplinary experience."}
    ]
}

# --- FUNCTIONS ---


def extract_text_pdf(file_path):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    doc = fitz.open(file_path)
    text = "".join([page.get_text() for page in doc])
    doc.close()
    return text

def extract_text_docx(file_path):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    doc = Document(file_path)
    content = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                content.append(cell.text)
    return "\n".join(content)

def detect_domain(filtered_tokens):
    text = " ".join(filtered_tokens)
    domain_scores = {}
    for domain, keywords in DOMAIN_KEYWORDS.items():
        matches = sum(1 for kw in keywords if re.search(r"\b" + re.escape(kw) + r"\b", text, re.IGNORECASE))
        if matches:
            domain_scores[domain] = matches
    return max(domain_scores, key=domain_scores.get) if domain_scores else "general"

def parse_resume_custom(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    text = extract_text_pdf(file_path) if ext == ".pdf" else extract_text_docx(file_path)
    tokens = word_tokenize(text.lower())
    stop_words = set(stopwords.words("english"))
    filtered_tokens = [w for w in tokens if w not in stop_words and len(w) > 2]

    domain = detect_domain(filtered_tokens)
    name_match = re.search(r"^([A-Z][a-z]+(?:\s[A-Z][a-z]+){0,2})", text, re.MULTILINE)
    name = name_match.group(1).strip() if name_match else "Not found"

    experience = re.findall(r"([a-zA-Z\s]+?(?:engineer|developer|intern|analyst|researcher|student|project)\s*(?:at\s+[\w\s]+)?)\s*(\d{4}(?:-\d{4}|-\s*present)?)", text, re.IGNORECASE)
    experience = [f"{role.strip()} ({period})" for role, period in experience[:5]]
    total_exp = len(experience)

    education_patterns = r"(bachelor|master|btech|mtech|phd|bs|ms)\s+(?:of|in)?\s*([a-zA-Z\s]+?)(?:\s+(?:from|at)\s+[\w\s]+)?"
    education = re.findall(education_patterns, text, re.IGNORECASE)
    education = [f"{deg.title()} in {field.title()}" for deg, field in education]

    return {"name": name, "domain": domain, "experience": experience, "education": education[:3], "total_experience": total_exp}

def generate_llm_advice(resume_data, goal):
    domain = resume_data.get("domain", "general")
    course_list = COURSE_LIBRARY.get(domain, COURSE_LIBRARY["general"])
    job_list = JOB_LIBRARY.get(domain, JOB_LIBRARY["general"])
    return {"course_suggestions": course_list[:3], "job_suggestions": job_list[:3]}

def fetch_job_listings(goal, api_key=None, location="India", page=1):
    if not api_key:
        return []
    url = f"https://jooble.org/api/{api_key}"
    params = {"keywords": goal, "location": location, "page": page, "per_page": 5}
    try:
        response = requests.get(url, params=params, timeout=10)
        if response.status_code == 200:
            data = response.json()
            return [{"title": j["title"], "company": j["company"], "location": j.get("location", location), "url": j["link"]} for j in data.get("jobs", [])[:5]]
    except Exception as e:
        print(f"Job API error: {e}")
    return []

def main():
    resume_file = input("Enter resume path (e.g., my_resume.pdf): ").strip()
    goal = input("Enter your career goal (e.g., 'become a data scientist'): ").strip()
    jooble_key = input("Enter Jooble API key (optional): ").strip()

    try:
        print("\nParsing resume...")
        resume_data = parse_resume_custom(resume_file)
        print(f"Name: {resume_data['name']}")
        print(f"Detected Domain: {resume_data['domain'].title()}")
        print(f"Experience: {resume_data['total_experience']} entries found.")
        print(f"Education: {', '.join(resume_data['education'])}\n")

        print("Generating career advice...\n")
        advice = generate_llm_advice(resume_data, goal)
        real_jobs = fetch_job_listings(goal, jooble_key)

        print("=== Recommended Online Courses ===")
        for i, c in enumerate(advice["course_suggestions"], 1):
            print(f"{i}. {c['course_name']} ({c['platform']})")
            print(f"   Why: {c['why_recommended']}")
            print(f"   Duration: {c['estimated_duration']}\n")

        print("=== Suggested Job Roles ===")
        all_jobs = advice["job_suggestions"] + real_jobs
        for i, j in enumerate(all_jobs, 1):
            title = j.get("job_title") or j.get("title")
            company = j.get("company") or j.get("company_type", "Various Companies")
            skills = ", ".join(j.get("required_skills", []))
            why = j.get("why_suitable", "Strong match for your background.")
            print(f"{i}. {title} at {company}")
            print(f"   Required Skills: {skills}")
            print(f"   Why Suitable: {why}")
            if "url" in j:
                print(f"   Apply: {j['url']}\n")
            else:
                print()

    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    main()
